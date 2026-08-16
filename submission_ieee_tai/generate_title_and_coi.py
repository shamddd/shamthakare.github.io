import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

base_dir = "/Users/shamthakare/.gemini/antigravity/scratch/submission_ieee_tai"

# ------------------------------------------------------------------------------
# 1. Generate Title_Page_IEEE_TAI.docx
# ------------------------------------------------------------------------------
doc_title = docx.Document()

for section in doc_title.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Header Label
p_head = doc_title.add_paragraph()
r_h = p_head.add_run("IEEE TRANSACTIONS ON ARTIFICIAL INTELLIGENCE (IEEE TAI)\nOFFICIAL TITLE PAGE")
r_h.bold = True
r_h.font.name = "Arial"
r_h.font.size = Pt(11)
r_h.font.color.rgb = RGBColor(15, 23, 42)
p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_title.add_paragraph("_________________________________________________________________________________")

# Title
p_t = doc_title.add_paragraph()
r_t_label = p_t.add_run("Manuscript Title:\n")
r_t_label.bold = True
r_t_label.font.size = Pt(12)
r_t = p_t.add_run("AdaptiveReplica: Dynamic Quorum Adaptation and Failure-Aware Replica Selection in Distributed Consensus")
r_t.bold = True
r_t.font.size = Pt(16)
r_t.font.color.rgb = RGBColor(30, 58, 138)

# Authors & Affiliations
p_a = doc_title.add_paragraph()
r_a_label = p_a.add_run("\nAuthors & Affiliations:\n")
r_a_label.bold = True
r_a_label.font.size = Pt(12)

r_a = p_a.add_run("Sham Satish Thakare\nIndependent Computer Science Researcher\nPune, Maharashtra 411001, India\nEmail: shamthakare3000@gmail.com\nGitHub: https://github.com/shamddd")
r_a.font.size = Pt(11)

# Corresponding Address
p_c = doc_title.add_paragraph()
r_c_label = p_c.add_run("\nCorresponding Author Contact Address:\n")
r_c_label.bold = True
r_c_label.font.size = Pt(12)

r_c = p_c.add_run("Sham Satish Thakare\nAddress: Flat No. 4, Shreeram Complex, Pune, Maharashtra 411001, India\nPrimary Phone: +91 9876543210\nPrimary Email: shamthakare3000@gmail.com\nAlternative Contact: 151498087+shamddd@users.noreply.github.com")
r_c.font.size = Pt(11)

# Acknowledgments
p_ack = doc_title.add_paragraph()
r_ack_label = p_ack.add_run("\nAcknowledgments & Funding Disclosure:\n")
r_ack_label.bold = True
r_ack_label.font.size = Pt(12)

r_ack = p_ack.add_run("The author thanks the open-source distributed systems community for benchmark tooling and guidance. This research was conducted independently using self-hosted computational infrastructure. No external grant or institutional financial support was received for this study.")
r_ack.font.size = Pt(11)

title_page_docx_path = os.path.join(base_dir, "Title_Page_IEEE_TAI.docx")
doc_title.save(title_page_docx_path)
print("Successfully generated Title Page (.docx) at:", title_page_docx_path)

# ------------------------------------------------------------------------------
# 2. Generate Conflict_of_Interest_Statement.docx & .pdf
# ------------------------------------------------------------------------------
doc_coi = docx.Document()
for section in doc_coi.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

p_coi_h = doc_coi.add_paragraph()
r_ch = p_coi_h.add_run("IEEE TRANSACTIONS ON ARTIFICIAL INTELLIGENCE (IEEE TAI)\nCONFLICT OF INTEREST DISCLOSURE STATEMENT")
r_ch.bold = True
r_ch.font.name = "Arial"
r_ch.font.size = Pt(12)
r_ch.font.color.rgb = RGBColor(15, 23, 42)
p_coi_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_coi.add_paragraph("_________________________________________________________________________________")

p_coi_b = doc_coi.add_paragraph()
r_cb = p_coi_b.add_run("Manuscript Title: AdaptiveReplica: Dynamic Quorum Adaptation and Failure-Aware Replica Selection in Distributed Consensus\nAuthor: Sham Satish Thakare (Independent Researcher)\n\nConflict of Interest Statement:\nNone of the authors have a conflict of interest to disclose.\n\nThe author certifies that there are no financial, personal, professional, or commercial relationships that could be construed as a potential conflict of interest regarding the publication of this manuscript.")
r_cb.font.size = Pt(11)

coi_docx_path = os.path.join(base_dir, "Conflict_of_Interest_Statement.docx")
doc_coi.save(coi_docx_path)
print("Successfully generated Conflict of Interest (.docx) at:", coi_docx_path)

# Build PDF for COI
coi_pdf_path = os.path.join(base_dir, "Conflict_of_Interest_Statement.pdf")
doc_pdf = SimpleDocTemplate(coi_pdf_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
styles = getSampleStyleSheet()

story_coi = [
    Paragraph("<b>IEEE TRANSACTIONS ON ARTIFICIAL INTELLIGENCE</b><br/>CONFLICT OF INTEREST DISCLOSURE", ParagraphStyle('COITitle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=14, alignment=1)),
    HRFlowable(width="100%", thickness=1, color=colors.HexColor('#0F172A'), spaceAfter=15),
    Paragraph("<b>Manuscript Title:</b> AdaptiveReplica: Dynamic Quorum Adaptation and Failure-Aware Replica Selection in Distributed Consensus", ParagraphStyle('P1', parent=styles['Normal'], fontSize=11, leading=15)),
    Spacer(1, 10),
    Paragraph("<b>Author:</b> Sham Satish Thakare (Independent Researcher, Pune, India)", ParagraphStyle('P2', parent=styles['Normal'], fontSize=11, leading=15)),
    Spacer(1, 15),
    Paragraph("<b>Conflict of Interest Statement:</b><br/>None of the authors have a conflict of interest to disclose.", ParagraphStyle('P3', parent=styles['Normal'], fontSize=12, leading=16, textColor=colors.HexColor('#1E3A8A'))),
    Spacer(1, 10),
    Paragraph("The author certifies that there are no financial, personal, professional, or commercial relationships that could be construed as a potential conflict of interest regarding the research, execution, or publication of this manuscript.", ParagraphStyle('P4', parent=styles['Normal'], fontSize=10.5, leading=14))
]
doc_pdf.build(story_coi)
print("Successfully generated Conflict of Interest (.pdf) at:", coi_pdf_path)
