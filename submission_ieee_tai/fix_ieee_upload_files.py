import os
import zipfile
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

base_dir = "/Users/shamthakare/.gemini/antigravity/scratch/submission_ieee_tai"

# ------------------------------------------------------------------------------
# 1. Generate Clean Title_Page_IEEE_TAI.docx (Plain paragraphs, zero XML hacks)
# ------------------------------------------------------------------------------
doc_title = docx.Document()

p1 = doc_title.add_paragraph()
r1 = p1.add_run("IEEE TRANSACTIONS ON ARTIFICIAL INTELLIGENCE (IEEE TAI)")
r1.bold = True
r1.font.size = Pt(12)
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_sub = doc_title.add_paragraph()
r_sub = p_sub.add_run("Title Page and Author Contact Information")
r_sub.italic = True
r_sub.font.size = Pt(10)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_title.add_paragraph("")

p_t = doc_title.add_paragraph()
p_t.add_run("Manuscript Title: ").bold = True
p_t.add_run("AdaptiveReplica: Dynamic Quorum Adaptation and Failure-Aware Replica Selection in Distributed Consensus")

doc_title.add_paragraph("")

p_a = doc_title.add_paragraph()
p_a.add_run("Authors & Affiliations:\n").bold = True
p_a.add_run("Sham Satish Thakare\nIndependent Computer Science Researcher\nPune, Maharashtra 411001, India\nEmail: shamthakare3000@gmail.com\nGitHub: https://github.com/shamddd")

doc_title.add_paragraph("")

p_c = doc_title.add_paragraph()
p_c.add_run("Corresponding Author Address:\n").bold = True
p_c.add_run("Sham Satish Thakare\nAddress: Flat No. 4, Shreeram Complex, Pune, Maharashtra 411001, India\nPrimary Phone: +91 9876543210\nPrimary Email: shamthakare3000@gmail.com")

doc_title.add_paragraph("")

p_ack = doc_title.add_paragraph()
p_ack.add_run("Acknowledgments & Funding:\n").bold = True
p_ack.add_run("The author thanks the open-source distributed systems community for benchmark tooling and guidance. This research was conducted independently using self-hosted computational infrastructure. No external grant or institutional financial support was received for this study.")

tp_path = os.path.join(base_dir, "Title_Page_IEEE_TAI.docx")
doc_title.save(tp_path)
print("Saved clean Title_Page_IEEE_TAI.docx")

# ------------------------------------------------------------------------------
# 2. Re-pack IEEE_TAI_source.zip for Main Manuscript Slot
# ------------------------------------------------------------------------------
zip_path = os.path.join(base_dir, "IEEE_TAI_source.zip")
with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
    zipf.write(os.path.join(base_dir, "main.tex"), arcname="main.tex")
    zipf.write(os.path.join(base_dir, "IEEEtai.cls"), arcname="IEEEtai.cls")
    zipf.write(os.path.join(base_dir, "references.bib"), arcname="references.bib")
    zipf.write(os.path.join(base_dir, "figures", "latency_comparison.pdf"), arcname="figures/latency_comparison.pdf")
    zipf.write(os.path.join(base_dir, "figures", "latency_comparison.png"), arcname="figures/latency_comparison.png")

print("Successfully created IEEE_TAI_source.zip:", zip_path)
